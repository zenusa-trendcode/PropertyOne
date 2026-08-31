from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent.parent
OUT = BASE_DIR / "docs" / "Zenusa PMS - User Guide.docx"

NAVY = RGBColor(30, 64, 175)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
TEAL = RGBColor(15, 118, 110)
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    dxa_widths = [int(width * 1440) for width in widths_in]
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in dxa_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(dxa_widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="CBD5E1"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_para(doc, text, size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, "BFDBFE")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    paragraph_text(p, title, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    paragraph_text(p2, body, size=10.5, color=INK)
    add_para(doc, "", after=4)


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], "E8EEF5")
        p = header_cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(text))
            set_run_font(run, size=9.5, color=INK)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_page_number(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def set_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        header.paragraphs[0].text = ""
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [3.25, 3.25])
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_text(left, "Zenusa PMS", size=9, color=MUTED, bold=True)
    paragraph_text(right, "User Guide", size=9, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        footer.paragraphs[0].text = ""
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [4.5, 2.0])
    fp_left = ftable.cell(0, 0).paragraphs[0]
    fp_right = ftable.cell(0, 1).paragraphs[0]
    fp_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_text(fp_left, "Confidential - Internal Use | Zenusa PMS", size=9, color=MUTED)
    add_page_number(fp_right)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc):
    add_para(doc, "PRODUCT OPERATIONS GUIDE", size=10.5, color=TEAL, bold=True, after=84, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    title = p.add_run("Zenusa PMS - User Guide")
    set_run_font(title, size=30, color=NAVY, bold=True)
    add_para(
        doc,
        "Panduan pengguna untuk operasional property management system",
        size=14,
        color=DARK_BLUE,
        after=30,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_callout(
        doc,
        "Dokumen ini untuk siapa?",
        "Property manager, finance AR, facility team, admin dokumen, owner representative, dan manajemen yang memantau portfolio properti.",
    )
    add_para(doc, "", after=96)
    add_data_table(
        doc,
        ["Version", "Date", "Prepared For"],
        [["1.0", "31 Agustus 2026", "Zenusa PMS users and implementation team"]],
        [1.1, 1.55, 3.85],
    )
    add_para(doc, "Confidential - for internal product demonstration and training use.", size=9.5, color=MUTED, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    add_data_table(
        doc,
        ["Item", "Detail"],
        [
            ["Document title", "Zenusa PMS - User Guide"],
            ["Application scope", "Portfolio, unit, tenant, lease, billing, maintenance, inspection, document, approval, and reporting workflows."],
            ["Primary users", "Property manager, finance AR, facility team, document admin, owner representative, and system admin."],
            ["Demo URL", "http://127.0.0.1:8000/"],
            ["Demo login", "Username: admin / Password: Admin12345!"],
        ],
        [1.7, 4.8],
    )
    add_callout(
        doc,
        "Quick start",
        "Login ke aplikasi, buka Command Center, pilih filter Property bila perlu, lalu lanjutkan ke modul yang sesuai dengan pekerjaan harian.",
    )


def add_navigation(doc):
    add_heading(doc, "2. Navigasi dan Command Center", 1)
    add_para(
        doc,
        "Command Center adalah halaman utama untuk melihat kondisi portfolio secara cepat. KPI yang ditampilkan mencakup Monthly Rent Roll, Occupancy, Outstanding AR, Open Work Orders, Deposit Held, dan Approvals.",
    )
    add_data_table(
        doc,
        ["Area", "Fungsi", "Tindakan pengguna"],
        [
            ["Portfolio Performance", "Memantau performa properti, okupansi, rent roll, open work order, dan status.", "Klik nama properti untuk membuka detail."],
            ["Cash Aging", "Melihat umur piutang berdasarkan bucket Current, 1-30, 31-60, 61-90, dan 90+.", "Prioritaskan invoice overdue dan escalated."],
            ["Priority Queue", "Menampilkan work order aktif dengan prioritas tinggi.", "Buka work order dan update status pekerjaan."],
            ["Lease Pipeline", "Melihat kontrak Active, Renewal Offered, Move Out Notice, dan Collection Hold.", "Follow up renewal atau move-out."],
            ["Documents Due", "Memantau dokumen pending, review, draft, atau mendekati expiry.", "Lengkapi metadata dan update status dokumen."],
            ["Approval Queue", "Menampilkan approval yang perlu keputusan.", "Review amount, object reference, dan approver."],
        ],
        [1.7, 2.45, 2.35],
    )


def add_module_matrix(doc):
    add_heading(doc, "3. Module Reference Matrix", 1)
    add_data_table(
        doc,
        ["Module", "Digunakan untuk", "Data kunci", "User utama"],
        [
            ["Organizations", "Mengelola legal entity, currency, tax, dan contact context.", "Name, legal name, tax ID, default currency.", "System admin"],
            ["Properties", "Master data properti dan konfigurasi operational ownership.", "Code, type, status, manager, service charge.", "Property manager"],
            ["Units", "Availability, kondisi unit, harga sewa, meter, dan next action.", "Unit code, status, market rent, current rent.", "Leasing / Property admin"],
            ["Tenants", "Identitas penyewa, KYC, risk, contact, dan outstanding balance.", "Tenant code, risk level, KYC, health score.", "Property manager / Finance"],
            ["Leases", "Kontrak sewa, lifecycle, rent, deposit, escalation, dan renewal.", "Lease code, dates, status, rent, deposit.", "Leasing / Legal"],
            ["Billing", "Invoice, payment, outstanding, reminder, dan aging AR.", "Invoice, due date, total, paid, aging.", "Finance AR"],
            ["Vendors", "Vendor maintenance, SLA, contact, rating, dan active status.", "Vendor code, category, SLA, rating.", "Facility lead"],
            ["Maintenance", "Work order queue, vendor assignment, cost, SLA, approval, dan evidence.", "WO number, priority, status, due date, cost.", "Facility team"],
            ["Inspections", "Jadwal inspeksi, compliance scope, score, finding, dan corrective action.", "Inspection no., scope, score, findings.", "Facility / HSE"],
            ["Documents", "Vault dokumen lease, KYC, handover, insurance, tax, dan compliance.", "Document no., type, status, expiry.", "Document admin"],
            ["Approvals", "Capex, discount, write-off, legal escalation, dan exception approval.", "Request no., amount, approver, status.", "Manager / Owner rep"],
            ["Reports", "Owner statement, delinquency pack, asset health, dan lease expiry book.", "Report pack and metrics.", "Owner / Management"],
        ],
        [1.15, 2.1, 1.9, 1.35],
    )


def add_daily_workflows(doc):
    add_heading(doc, "4. Workflow Harian per Role", 1)
    workflows = [
        ("Property Manager", [
            "Buka Command Center dan filter property yang menjadi tanggung jawab.",
            "Cek occupancy, open work order, approval pending, dan documents due.",
            "Buka Lease Pipeline untuk kontrak renewal atau move-out.",
            "Follow up tenant dengan risk Monitor atau At Risk.",
        ]),
        ("Finance AR", [
            "Buka Billing dan filter invoice overdue atau escalated.",
            "Cek aging bucket, outstanding amount, dan reminder count.",
            "Catat janji bayar di notes dan update payment jika sudah diterima.",
            "Buat approval untuk legal escalation atau write-off bila diperlukan.",
        ]),
        ("Facility Team", [
            "Buka Maintenance dan prioritaskan Emergency serta High.",
            "Assign vendor sesuai SLA dan kategori pekerjaan.",
            "Update status dari triage sampai QA Review dan Closed.",
            "Pastikan before/after evidence tersedia sebelum work order ditutup.",
        ]),
        ("Owner Representative", [
            "Cek rent roll, occupancy, outstanding, deposit held, dan approval queue.",
            "Review approval bernilai besar atau berdampak legal.",
            "Gunakan Reports untuk Owner Statement dan Asset Health.",
            "Pantau property Watchlist dan action plan manager.",
        ]),
    ]
    for role, steps in workflows:
        add_heading(doc, role, 2)
        add_numbers(doc, steps)


def add_core_guides(doc):
    add_heading(doc, "5. Panduan Modul Operasional", 1)
    sections = [
        (
            "Property dan Unit Setup",
            "Gunakan Properties untuk membuat master properti, lalu Units untuk menyiapkan inventory unit yang bisa disewakan.",
            ["Setiap property wajib punya kode unik.", "Setiap unit harus unik dalam satu property.", "Isi market rent sebagai acuan pricing dan current rent saat unit aktif disewa.", "Gunakan next action untuk pekerjaan harian seperti Publish listing, Schedule handover, atau Payment plan."],
        ),
        (
            "Tenant dan Lease Lifecycle",
            "Tenant menyimpan identitas dan risiko penyewa; Lease mengikat tenant, unit, property, rent, deposit, dan masa kontrak.",
            ["KYC harus dipantau sampai Complete.", "Lease Draft dipakai saat kontrak masih disiapkan.", "Renewal Offered dipakai saat penawaran renewal sudah dikirim.", "Move Out Notice harus diikuti dokumen handover dan work order make-ready.", "Collection Hold dipakai untuk lease yang terdampak tunggakan."],
        ),
        (
            "Billing dan Collection",
            "Billing mengelola invoice, status pembayaran, aging bucket, reminder, dan outstanding amount.",
            ["Issued berarti invoice sudah diterbitkan.", "Partially Paid berarti ada pembayaran sebagian.", "Overdue harus masuk follow-up collection.", "Escalated dipakai untuk tindakan legal atau eskalasi manajemen.", "Gunakan reminder count dan notes untuk rekam jejak collection."],
        ),
        (
            "Maintenance dan Vendor Control",
            "Maintenance dipakai untuk work order, prioritas, vendor assignment, SLA, budget approval, dan evidence rule.",
            ["Emergency harus punya due date dalam hitungan jam.", "Biaya besar perlu requires owner approval.", "Estimated cost, approved budget, dan actual cost harus dipantau.", "Gunakan quick transition: Vendor Assigned, In Progress, QA Review, Closed.", "Jangan close work order sebelum evidence lengkap."],
        ),
        (
            "Inspection, Documents, dan Approvals",
            "Inspection menjaga compliance dan asset condition; Documents menjaga metadata dokumen; Approvals menjaga kontrol keputusan bisnis.",
            ["Inspection score rendah harus menghasilkan corrective action.", "Dokumen compliance harus punya expiry date.", "Approval dibuat untuk capex, discount, write-off, legal collection notice, atau lease exception.", "Decision note harus diisi setelah approved atau rejected."],
        ),
    ]
    for title, intro, bullets in sections:
        add_heading(doc, title, 2)
        add_para(doc, intro)
        add_bullets(doc, bullets)


def add_status_reference(doc):
    add_heading(doc, "6. Status dan Aturan Bisnis", 1)
    add_data_table(
        doc,
        ["Object", "Status penting", "Kapan digunakan"],
        [
            ["Property", "Active, Stabilized, Lease Up, Watchlist, Inactive", "Menandai fase operasional properti."],
            ["Unit", "Vacant, Make Ready, Occupied, Notice, Delinquent, Offline", "Menandai availability dan risiko unit."],
            ["Tenant", "Excellent, Monitor, At Risk", "Menandai risiko tenant berdasarkan pembayaran, dokumen, dan komunikasi."],
            ["Lease", "Draft, Active, Renewal Offered, Move Out Notice, Collection Hold", "Menandai lifecycle kontrak sewa."],
            ["Invoice", "Draft, Issued, Partially Paid, Paid, Overdue, Escalated, Void", "Menandai collection dan pembayaran."],
            ["Work Order", "New, Triaged, Waiting Approval, Vendor Assigned, Waiting Parts, In Progress, QA Review, Closed", "Menandai progres pekerjaan maintenance."],
            ["Inspection", "Scheduled, In Progress, Needs Action, Critical, Complete", "Menandai status pemeriksaan dan compliance."],
            ["Document", "Draft, Pending, Review, Complete, Expired", "Menandai kelengkapan dan masa berlaku dokumen."],
            ["Approval", "Pending, Approved, Rejected, Cancelled", "Menandai keputusan manager atau owner."],
        ],
        [1.2, 2.65, 2.65],
    )
    add_heading(doc, "Data governance rules", 2)
    add_bullets(
        doc,
        [
            "Lease harus menghubungkan property, unit, dan tenant yang benar.",
            "Invoice sebaiknya selalu terkait tenant dan property.",
            "Work order emergency wajib punya due date jelas.",
            "Approval wajib dibuat untuk biaya besar, diskon besar, write-off, atau tindakan legal.",
            "Tenant risk harus diperbarui jika ada tunggakan atau dokumen bermasalah.",
        ],
    )


def add_reports_and_admin(doc):
    add_heading(doc, "7. Reports dan Admin", 1)
    add_data_table(
        doc,
        ["Report pack", "Isi utama", "Pengguna"],
        [
            ["Owner Statement", "Rent roll, NOI, deposit held, owner payout, dan capex reserve.", "Owner representative"],
            ["Delinquency Pack", "Outstanding, aging bucket, reminder history, promise-to-pay, dan legal escalation.", "Finance AR"],
            ["Asset Health", "Open WO, SLA breach, inspection score, recurring fault, dan vendor rating.", "Facility lead"],
            ["Lease Expiry Book", "Lease expiry 30/60/90 hari, renewal probability, market rent gap, dan make-ready date.", "Property manager"],
        ],
        [1.55, 3.45, 1.5],
    )
    add_callout(
        doc,
        "Django Admin",
        "Admin tersedia di http://127.0.0.1:8000/admin/ untuk mengelola model data, inspection finding, user internal, dan debugging data selama implementasi awal.",
    )


def add_limitations(doc):
    add_heading(doc, "8. Batasan Versi Saat Ini dan Roadmap", 1)
    add_para(
        doc,
        "Versi saat ini sudah berjalan sebagai aplikasi Django dengan database, CRUD inti, dashboard, seed demo, admin, dan workflow work order. Untuk dijual luas, fitur berikut perlu diperkuat sebagai roadmap komersial.",
    )
    add_bullets(
        doc,
        [
            "Role-based access control per modul dan per property.",
            "Multi-company tenancy untuk banyak client.",
            "Upload file dokumen asli dengan object storage.",
            "Export PDF/XLSX untuk owner statement dan finance report.",
            "Email atau WhatsApp reminder otomatis.",
            "Payment gateway atau bank statement matching.",
            "Audit trail detail per perubahan field.",
            "Owner portal, tenant portal, dan API untuk mobile app.",
            "Production deployment dengan PostgreSQL, backup, monitoring, dan hardening security.",
        ],
    )


def add_troubleshooting(doc):
    add_heading(doc, "9. Troubleshooting", 1)
    add_data_table(
        doc,
        ["Masalah", "Langkah pengecekan"],
        [
            ["Aplikasi tidak terbuka", "Jalankan .\\.venv\\Scripts\\python.exe manage.py runserver 127.0.0.1:8000"],
            ["Database belum siap", "Jalankan .\\.venv\\Scripts\\python.exe manage.py migrate"],
            ["Data demo belum ada", "Jalankan .\\.venv\\Scripts\\python.exe manage.py seed_demo"],
            ["Login admin gagal", "Jalankan ulang seed_demo untuk reset password admin."],
            ["Ingin cek kesehatan sistem", "Jalankan manage.py check dan manage.py test."],
        ],
        [2.0, 4.5],
    )
    add_para(doc, "End of document.", size=9.5, color=MUTED, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)


def build():
    doc = Document()
    setup_styles(doc)
    props = doc.core_properties
    props.title = "Zenusa PMS - User Guide"
    props.subject = "Property Management System user guide"
    props.author = "Zenusa PMS"
    props.comments = "Generated from PropertiOne PMS user guide source."
    props.created = date.today()

    add_cover(doc)
    add_document_control(doc)
    add_navigation(doc)
    add_module_matrix(doc)
    add_daily_workflows(doc)
    add_core_guides(doc)
    add_status_reference(doc)
    add_reports_and_admin(doc)
    add_limitations(doc)
    add_troubleshooting(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
